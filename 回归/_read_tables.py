import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r'd:/HuaweiMoveData/Users/27438/Desktop/回归/论文完整版.docx')
print(f'表格数: {len(doc.tables)}\n')
for i, t in enumerate(doc.tables):
    print(f'=== 表格 {i+1} ({len(t.rows)}行 x {len(t.columns)}列) ===')
    for j, r in enumerate(t.rows):
        cells = [c.text.strip().replace('\n',' ') for c in r.cells]
        print(' | '.join(cells))
    print()

# 关键段落：第五、六章
for pi, p in enumerate(doc.paragraphs):
    if 'Heading 1' in p.style.name or 'Heading 2' in p.style.name:
        print(f'[{pi}] {p.style.name}: {p.text}')
    elif '实证结果' in p.text or '基准回归' in p.text or '描述性统' in p.text and pi > 60:
        print(f'[{pi}] {p.text[:200]}')
